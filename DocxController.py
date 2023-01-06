import os
import re
from docx import Document
from docx.shared import Inches, Pt
import ProgressBarController
class DocxController:
    __nElement = None
    __doc = None
    __title = None
    def __init__(self,title):
        self.__doc = Document()
        self.__title = re.sub(".pdf",".docx",title)
        #ADD only the name (not the path) without .docx extension
        self.__doc.add_heading(os.path.basename(re.sub(".docx","",self.__title)),level = 0)
        

    def InsertText(self,text):
        p = self.__doc.add_paragraph().add_run(text)
        #CALIBRI 12
        p.font.size = Pt(12)
        p.font.name = 'Calibri'
    def InsertPhoto(self,imageName):
        self.__doc.add_picture(imageName, width = Inches(6.61), height = Inches(3.9))
        os.remove("./"+imageName)

#optional parameters are required for the case where it is not batch
    def Save(self, batch = False, folder = None, nElements = None):
        if batch:
            if not os.path.isdir(folder+"/"+"DOCX"):
                os.mkdir(folder+"/"+"DOCX")
            self.__doc.save(folder + "/" + "DOCX/" + self.__title.replace(folder,""))
            ProgressBarController.updateValue(int(100/nElements))
        else:
            self.__doc.save(self.__title)
            ProgressBarController.updateValue(100)
